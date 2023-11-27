from inference.engine import Model
from pdf2docx import parse
import copy
import streamlit as st
import mammoth
import docx
from io import BytesIO
import timeit
import time
import streamlit.components.v1 as components
import base64

from docx import Document

def download_button(object_to_download, download_filename):
    """
    Generates a link to download the given object_to_download.
    Params:
    ------
    object_to_download:  The object to be downloaded.
    download_filename (str): filename and extension of file. e.g. mydata.docx,
    Returns:
    -------
    (str): the anchor tag to download object_to_download
    """
    if isinstance(object_to_download, list):
        doc = Document()
        for line in object_to_download:
            doc.add_paragraph(line)
        doc.save(download_filename)
        with open(download_filename, "rb") as doc_file:
            object_to_download = doc_file.read()

    try:
        # some strings <-> bytes conversions necessary here
        b64 = base64.b64encode(object_to_download).decode()

    except AttributeError as e:
        b64 = base64.b64encode(object_to_download).decode()

    dl_link = f"""
    <html>
    <head>
    <title>Start Auto Download file</title>
    <script src="http://code.jquery.com/jquery-3.2.1.min.js"></script>
    <script>
    $('<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{download_filename}">')[0].click()
    </script>
    </head>
    </html>
    """
    return dl_link


def generate_docx(array_content):
    # Create a DOCX file from the array content
    doc = docx.Document()
    for line in array_content:
        doc.add_paragraph(line)

    # Save the document to a BytesIO object
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer

def extract_text_from_docx(uploaded_file):
    result = mammoth.extract_raw_text(uploaded_file)
    text = result.value.splitlines()
    # Filter out empty lines
    text = [line for line in text if line.strip()]
    return text

def save_text_to_docx(text, file_path):
    doc = docx.Document()
    for line in text:
        doc.add_paragraph(line)
    
    # Save the document to a BytesIO object
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def translate_sentences(sentences):
    print("====== translating =======")
    src_lang = "hin_Deva"
    tgt_lang = "eng_Latn"
    model = Model("./ct2_model/indic-en-deploy/ct2_int8_model", device="cpu", model_type="ctranslate2")
    # Perform translation using your inference engine model
    # Assuming 'batch_translate' is a function that performs batch translation
    en_translations = model.batch_translate(sentences, src_lang, tgt_lang)
    return en_translations

def main():
    st.title("DOCX Text Translation")

    uploaded_file = st.file_uploader("Upload DOCX or PDF file", type=["docx", "pdf"])
    if uploaded_file is not None:
        pdf_file_path = "temp.pdf"
        docx_file_path = "temp.docx"
        # Check the file type of the uploaded file
        if uploaded_file.type == "application/pdf":
            # Convert the PDF file to DOCX format
            with open(pdf_file_path, "wb") as f:
                f.write(uploaded_file.getvalue())
            parse(pdf_file_path, docx_file_path)
            uploaded_file = open(docx_file_path, "rb")

        original_text = extract_text_from_docx(uploaded_file)

        print(st.session_state)
        # Check if translations already exist in the session state
        if 'translation_done' not in st.session_state or not st.session_state.translation_done: 
            start_time = time.time()
            translated_text = translate_sentences(original_text)
            end_time = time.time()
            execution_time = end_time - start_time
            print(f"Execution time: {execution_time} seconds")
            # Save the translations to the session state
            st.session_state.translated_text = translated_text
            st.session_state.translation_done = True
        else:
            # If translations already exist in the session state, use them
            translated_text = st.session_state.translated_text

        
        # Collect translations
        for index, paragraph in enumerate(original_text):
            st.write(f"Line {index + 1}")
            st.write(f"Original Text: {paragraph}")
            translation = st.text_area(f"Translated Line {index + 1}", value=translated_text[index])
            translated_text[index] = translation  # Update translation value
            
        def download_df():
            sentences = translated_text 
            components.html(
                download_button(sentences, "result.docx"),
                height=0,
            )

        st.button("Download", on_click=download_df)


if __name__ == '__main__':
    main()
